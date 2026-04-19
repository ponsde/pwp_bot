"""Convert paper/paper.md to paper/赛题名称.docx.

Keeps paragraphs / headings / tables / code blocks / inline bold. The output is
intentionally plain: no per-element font styling beyond heading levels and
monospace for code blocks, so the user can adjust typography in Word.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


SRC = Path("/home/ponsde/taidi_bei/paper/paper.md")
DST = Path("/home/ponsde/taidi_bei/paper/上市公司财报智能问数助手.docx")


# ---------------------------------------------------------------------------
# Inline parsing
# ---------------------------------------------------------------------------

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_CODE_RE = re.compile(r"`([^`]+?)`")


def add_inline(paragraph, text: str) -> None:
    """Render a line with **bold** and `code` spans into runs."""
    i = 0
    while i < len(text):
        b = _BOLD_RE.search(text, i)
        c = _CODE_RE.search(text, i)
        # Pick the nearest marker
        candidates = [m for m in (b, c) if m is not None]
        if not candidates:
            paragraph.add_run(text[i:])
            return
        nxt = min(candidates, key=lambda m: m.start())
        if nxt.start() > i:
            paragraph.add_run(text[i : nxt.start()])
        run = paragraph.add_run(nxt.group(1))
        if nxt is b:
            run.bold = True
        else:
            run.font.name = "Consolas"
            # Ensure CJK font too, so Chinese chars inside code don't look odd
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        i = nxt.end()


# ---------------------------------------------------------------------------
# Block walker
# ---------------------------------------------------------------------------

def _is_table_header(lines: list[str], i: int) -> bool:
    if i + 1 >= len(lines):
        return False
    head = lines[i]
    sep = lines[i + 1]
    if not (head.strip().startswith("|") and head.strip().endswith("|")):
        return False
    # Separator row: | --- | :--- | ---: |
    cells = [c.strip() for c in sep.strip().strip("|").split("|")]
    return all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c)


def _split_table_row(row: str) -> list[str]:
    return [c.strip() for c in row.strip().strip("|").split("|")]


def main() -> int:
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Document defaults: enlarge base font, set CJK font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Page margins: standard A4 with 2.5cm margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(md):
        line = md[i]
        stripped = line.rstrip()

        # Code fence
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(9)
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Horizontal rule
        if stripped in {"---", "***"}:
            doc.add_paragraph("").add_run()
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # Tables
        if _is_table_header(md, i):
            header = _split_table_row(md[i])
            i += 2  # skip header + separator
            rows = []
            while i < len(md) and md[i].strip().startswith("|"):
                rows.append(_split_table_row(md[i]))
                i += 1
            tbl = doc.add_table(rows=len(rows) + 1, cols=len(header))
            tbl.style = "Light Grid Accent 1"
            for c, txt in enumerate(header):
                cell = tbl.rows[0].cells[c]
                cell.text = ""
                add_inline(cell.paragraphs[0], txt)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for r, row in enumerate(rows, start=1):
                for c, txt in enumerate(row):
                    cell = tbl.rows[r].cells[c] if c < len(tbl.rows[r].cells) else None
                    if cell is None:
                        continue
                    cell.text = ""
                    add_inline(cell.paragraphs[0], txt)
            doc.add_paragraph()  # spacing after table
            continue

        # Blockquote
        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped[2:])
            i += 1
            continue

        # Bullet / numbered list
        m_b = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        m_n = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m_b:
            indent = len(m_b.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            if indent:
                p.paragraph_format.left_indent = Cm(0.75 * indent)
            add_inline(p, m_b.group(2))
            i += 1
            continue
        if m_n:
            indent = len(m_n.group(1)) // 2
            p = doc.add_paragraph(style="List Number")
            if indent:
                p.paragraph_format.left_indent = Cm(0.75 * indent)
            add_inline(p, m_n.group(3))
            i += 1
            continue

        # Blank → paragraph break
        if not stripped:
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(str(DST))
    print(f"wrote {DST} ({DST.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
