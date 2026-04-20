"""Insert the four Web UI screenshots into the 5.4 section of the paper docx.

Strategy: open the regenerated docx, walk paragraphs until we hit the
"5.4 Web UI 与容器化" heading, then right before the next top-level
heading we splice in four (image + caption) pairs.

Run after ``scripts/md2docx.py``.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/home/ponsde/taidi_bei")
DOCX = ROOT / "paper" / "上市公司财报智能问数助手.docx"
FIGURES = [
    {
        "path": ROOT / "Screenshot 2026-04-20 082227.png",
        "caption": "图 5-1  Web UI 首页看板 —— 上市公司数、财报记录、研报资源、"
                   "最新报告期四项核心指标；下方为 System Health 与 Resource Stats。",
    },
    {
        "path": ROOT / "Screenshot 2026-04-20 084009.png",
        "caption": "图 5-2  Text2SQL 基础问答 —— 自然语言问题触发 mcp_fin_query "
                   "工具调用，展开后可见生成的 SQL、SQLite 返回行，"
                   "最终 bot 话术整合为简洁回答。",
    },
    {
        "path": ROOT / "Screenshot 2026-04-20 151642.png",
        "caption": "图 5-3  自动图表生成 —— 对趋势类问题 mcp_fin_query 自动渲染 "
                   "折线图嵌入回答；工具块顶部显示 “3 条引用” 与 “图表” 徽章。",
    },
    {
        "path": ROOT / "Screenshot 2026-04-20 151057.png",
        "caption": "图 5-4  研报 RAG 归因问答 —— openviking_search 返回 10 条 "
                   "参考文献，以编号卡片列出，配合附件 7 表 5 的 references "
                   "字段规范。",
    },
]


def _find_54_section(doc):
    """Return (start_index, end_index) paragraphs indexes bounding the 5.4 body.

    start_index: the heading paragraph itself
    end_index:   the next H2 heading after 5.4 (exclusive)
    """
    body = list(doc.paragraphs)
    start = None
    for i, p in enumerate(body):
        if p.style.name.startswith("Heading") and "Web UI" in p.text:
            start = i
            break
    if start is None:
        return None, None
    end = len(body)
    for j in range(start + 1, len(body)):
        if body[j].style.name.startswith("Heading 2"):
            end = j
            break
    return start, end


def _insert_paragraph_after(doc, anchor_el, text: str = "", style: str | None = None):
    """Insert a new paragraph in the document right after anchor_el.

    docx XML trick: python-docx doesn't have a built-in "insert after" for
    Paragraph objects at the document level, so we splice an OXML element.
    """
    from docx.oxml import OxmlElement

    new_p = OxmlElement("w:p")
    anchor_el.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, doc.paragraphs[0]._parent)
    if text:
        run = p.add_run(text)
        run.font.size = Pt(9)
    if style:
        try:
            p.style = doc.styles[style]
        except KeyError:
            pass
    return p


def main() -> int:
    if not DOCX.exists():
        print(f"missing {DOCX}; run scripts/md2docx.py first", file=sys.stderr)
        return 1

    doc = Document(str(DOCX))
    start, end = _find_54_section(doc)
    if start is None:
        print("could not find 5.4 section heading", file=sys.stderr)
        return 2

    # Figures are inserted right before the paragraph at `end` — i.e. at the
    # tail of section 5.4, in order. We walk backwards so successive
    # insertions don't shift indices.
    #
    # The anchor we append after is the LAST body paragraph of 5.4 (index
    # end - 1 in the original list). Each "figure pair" = image paragraph +
    # caption paragraph, both inserted at the tail.
    last_idx = end - 1
    anchor_p = doc.paragraphs[last_idx]

    for fig in FIGURES:
        path = fig["path"]
        if not path.is_file():
            print(f"  skip (missing): {path.name}")
            continue
        # Add image paragraph
        img_para = _insert_paragraph_after(doc, anchor_p._element)
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(str(path), width=Cm(14))
        # Caption directly below
        cap_para = _insert_paragraph_after(doc, img_para._element, text=fig["caption"])
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor_p = cap_para  # so the next figure goes below this caption
        print(f"  inserted: {path.name}")

    doc.save(str(DOCX))
    print(f"wrote {DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
